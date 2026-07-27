#!/usr/bin/env python3
"""
Telecom Watch — 국제비교 소스 자동 검증
매주 월요일 실행: 각 국제 비교 보고서(OECD·ITU·총무성 등)의 신판 발표 여부를
Claude 웹 검색으로 확인하고, 새 발표가 있으면 intl.json 갱신 + 팀용 요약 보고서 생성.

수동 실행: INTL_CHECK=1 python scripts/intl_check.py
"""
from __future__ import annotations
import os, json, time
from datetime import datetime, timezone, timedelta

import anthropic

KST = timezone(timedelta(hours=9))
TODAY = datetime.now(KST).strftime('%Y-%m-%d')
INTL_FILE = 'data/intl.json'
REPORTS_FILE = 'data/intl_reports.json'
MAX_REPORTS = 30   # 보고서 최대 보관 수

client = anthropic.Anthropic(api_key=os.environ['ANTHROPIC_API_KEY'])


def report_signature(source_id: str, latest: str) -> str:
    """소스ID + 최신판 표기 속 연도로 서명 생성 (표현 변화에도 안정적인 중복 판별용)."""
    import re as _re
    years = ''.join(sorted(set(_re.findall(r'(?:19|20)\d{2}', latest or ''))))
    return f'{source_id}|{years}'


def check_source(item: dict) -> dict | None:
    """웹 검색으로 해당 소스의 신판 발표 여부 확인. 새 발표 시 dict 반환."""
    prompt = f"""당신은 한국 이동통신사(SKT) 서비스제도팀의 리서치 어시스턴트입니다.
아래 국제 통신요금/시장 비교 보고서의 최신 발표 현황을 웹 검색으로 확인해주세요.

[확인 대상]
- 발행기관: {item.get('org','')}
- 보고서명: {item.get('name','')}
- 발표 주기: {item.get('cycle','')}
- 현재 기록된 최신판: {item.get('latest','')}

[할 일]
1. 웹 검색으로 이 보고서의 가장 최근 발표(판/연도)를 확인
2. 현재 기록된 최신판보다 새로운 발표가 있는지 판단
3. 새 발표가 있으면 한국 관련 핵심 결과를 파악

[응답 형식 — 아래 JSON만 출력, 다른 텍스트 금지]
{{
  "new_edition": true 또는 false,
  "latest": "확인된 최신판 표기 (예: 2026년판 (2026.5 공표))",
  "result": "한국 관련 핵심 결과 1~2문장 (새 발표 없으면 빈 문자열)",
  "team_report": "팀 보고용 요약 3~4문장: 무엇이 발표됐고, 한국 순위/평가가 어떻게 나왔고, 대응 관점에서 뭘 봐야 하는지 (새 발표 없으면 빈 문자열)",
  "source_url": "확인에 사용한 대표 출처 URL (없으면 빈 문자열)"
}}

확실하지 않으면 new_edition은 false로 판단하세요. 추측으로 새 발표를 만들어내지 마세요."""

    try:
        msg = client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=1024,
            tools=[{
                'type': 'web_search_20250305',
                'name': 'web_search',
                'max_uses': 2,
            }],
            messages=[{'role': 'user', 'content': prompt}],
        )
        # 최종 텍스트 블록에서 JSON 추출
        text = ''
        for block in msg.content:
            if block.type == 'text':
                text += block.text
        text = text.strip()
        if '```' in text:
            text = text.split('```')[1]
            if text.startswith('json'):
                text = text[4:]
        start = text.find('{')
        end = text.rfind('}')
        if start < 0 or end < 0:
            return None
        data = json.loads(text[start:end + 1])
        return data
    except Exception as e:
        print(f'  ⚠️  검증 오류 ({item.get("id","")}): {e}')
        return None


def intl_check():
    if not os.path.exists(INTL_FILE):
        print('⏭  intl.json 없음 — 건너뜀')
        return

    with open(INTL_FILE, 'r', encoding='utf-8') as f:
        intl = json.load(f)
    items = intl.get('items', [])

    reports = []
    if os.path.exists(REPORTS_FILE):
        try:
            with open(REPORTS_FILE, 'r', encoding='utf-8') as f:
                reports = json.load(f)
        except Exception:
            reports = []

    # 기존 보고서로부터 이미 통보한 (소스+연도) 서명 집합 구성
    reported_sigs = {report_signature(r.get('source_id', ''), r.get('latest', ''))
                     for r in reports}

    print(f'🌐 국제비교 소스 자동 검증 시작 ({len(items)}개 소스)')
    updated = 0
    checked = 0
    skipped = 0
    from datetime import date
    today_d = datetime.now(KST).date()
    for item in items:
        # 최근 5일 내 검증한 소스는 스킵 (수동 재실행 시 중복 비용 방지)
        v = item.get('verified', '')
        try:
            if v and (today_d - date.fromisoformat(v)).days < 5:
                skipped += 1
                continue
        except Exception:
            pass
        print(f'  🔍 {item.get("name","")[:50]}')
        res = check_source(item)
        time.sleep(1)
        if res is not None:
            item['verified'] = TODAY   # 검증 수행일 스탬프 (신판 여부 무관)
            checked += 1
        if not res or not res.get('new_edition'):
            continue

        new_latest = res.get('latest', '')
        if not new_latest or new_latest == item.get('latest'):
            continue

        # 재발송 방지: 같은 소스의 같은 연도 판을 이미 보고했으면 건너뜀
        # (AI가 표현만 살짝 바꿔 '신판'으로 재감지하는 경우 방지)
        sig = report_signature(item.get('id', ''), new_latest)
        if sig in reported_sigs:
            print(f'  ↩︎ 이미 보고된 판 — 재발송 생략: {new_latest}')
            item['latest'] = new_latest   # 표기만 갱신, 알림은 생략
            continue
        reported_sigs.add(sig)

        print(f'  🆕 신판 확인: {new_latest}')
        item['latest'] = new_latest
        if res.get('result'):
            item['result'] = res['result']
        updated += 1

        if res.get('team_report'):
            reports.insert(0, {
                'date': TODAY,
                'source_id': item.get('id', ''),
                'org': item.get('org', ''),
                'name': item.get('name', ''),
                'latest': new_latest,
                'report': res['team_report'],
                'url': res.get('source_url', '') or item.get('url', ''),
            })

    if checked > 0:
        intl['updated'] = TODAY
        with open(INTL_FILE, 'w', encoding='utf-8') as f:
            json.dump(intl, f, ensure_ascii=False, indent=2)
    if updated > 0:
        with open(REPORTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(reports[:MAX_REPORTS], f, ensure_ascii=False, indent=2)
        print(f'✅ 신판 {updated}건 반영 + 팀 보고서 생성 (검증 {checked}건)')
        notify_new_editions(reports[:updated])
    else:
        print(f'  신판 없음 — 검증 {checked}건 확인일 갱신 (최근 검증 스킵 {skipped}건)')


def make_report_doc(r: dict) -> str:
    """팀 보고서를 정식 워드(.docx) 파일로 생성해 경로 반환 (폰·PC 모두 열림)."""
    import sys, os as _os
    sys.path.insert(0, _os.path.dirname(_os.path.abspath(__file__)))
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        from collect import tg_clean_body as _clean
    except Exception:
        def _clean(t): return str(t or '').strip()

    ACCENT = RGBColor(0x36, 0x17, 0xCE)
    doc = Document()
    # 기본 글꼴
    normal = doc.styles['Normal']
    normal.font.name = '맑은 고딕'
    normal.font.size = Pt(11)

    title = doc.add_heading('발표 요약 보고서', level=0)
    for run in title.runs:
        run.font.color.rgb = ACCENT

    meta = doc.add_paragraph()
    mrun = meta.add_run(f'작성일: {TODAY}  ·  작성: Telecom Watch (자동 생성)')
    mrun.font.size = Pt(9)
    mrun.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # 기본 정보 표
    rows = [('발행기관', r.get('org')), ('보고서명', r.get('name')),
            ('최신판', _clean(r.get('latest'))), ('원문', r.get('url'))]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v or '')

    h = doc.add_heading('1. 발표 내용 및 시사점', level=1)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    doc.add_paragraph(_clean(r.get('report')))

    foot = doc.add_paragraph()
    frun = foot.add_run('본 문서는 AI가 자동 생성한 참고 자료입니다. 대외 인용 전 원문 확인이 필요합니다.')
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    import re as _re
    safe = _re.sub(r'[\\/:*?"<>| ]', '_', r.get('org', 'report'))
    path = f'/tmp/발표요약보고서_{safe}_{TODAY}.docx'
    doc.save(path)
    return path


def notify_new_editions(new_reports: list):
    """신판 발견 즉시 텔레그램 통보 (보고서 전문 + 워드 파일 첨부)."""
    import sys
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    try:
        from collect import tg_send, tg_send_document, tg_chat_ids, tg_clean_body, tg_link
    except Exception as ex:
        print(f'  ⚠️  텔레그램 헬퍼 로드 실패: {ex}')
        return
    if not os.environ.get('TELEGRAM_BOT_TOKEN', '') or not tg_chat_ids():
        return
    for r in new_reports:
        link = tg_link('원문 보기', r.get('url', ''))
        msg = (f"🆕 *국제비교 신판 발표 감지*\n\n"
               f"*{r.get('org','')} — {r.get('name','')}*\n"
               f"최신판: {tg_clean_body(r.get('latest',''))}\n\n"
               f"{tg_clean_body(r.get('report',''))}"
               + (f"\n\n{link}" if link else ""))
        tg_send(msg)
        try:
            path = make_report_doc(r)
            tg_send_document(path, f"{r.get('org','')} 발표 요약 보고서")
        except Exception as ex:
            print(f'  ⚠️  워드 첨부 실패: {ex}')
    print(f'✅ 신판 통보 발송 ({len(new_reports)}건)')


def check_schedules():
    """국회 과방위·방통위 등 규제기관의 향후 일정을 웹 검색으로 확인해 calendar.json에 병합."""
    prompt = (f"오늘 날짜는 {TODAY}입니다. 웹 검색으로 다음을 확인해주세요:\n"
        "1. 국회 과학기술정보방송통신위원회(과방위)의 향후 예정된 전체회의·법안소위 일정\n"
        "2. 방송통신위원회(방통위) 전체회의 일정\n"
        "3. 과기정통부의 통신 관련 공청회·의견수렴 일정\n\n"
        '확인된 미래 일정만 JSON으로 출력: {"events": [{"date": "YYYY-MM-DD", "title": "일정 제목", "org": "기관명"}]}\n'
        "확실한 날짜가 확인된 것만 포함. 없으면 빈 배열.")
    try:
        msg = client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=800,
            tools=[{'type': 'web_search_20250305', 'name': 'web_search', 'max_uses': 3}],
            messages=[{'role': 'user', 'content': prompt}],
        )
        text = ''
        for block in msg.content:
            if block.type == 'text':
                text += block.text
        text = text.strip()
        if '```' in text:
            text = text.split('```')[1]
            if text.startswith('json'):
                text = text[4:]
        start = text.find('{')
        end = text.rfind('}')
        events = json.loads(text[start:end+1]).get('events', []) if start >= 0 else []
    except Exception as ex:
        print(f'  ⚠️  일정 검색 오류: {ex}')
        return
    # collect.py의 병합 로직 재사용
    import sys
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    try:
        from collect import merge_calendar_events
        merge_calendar_events(events, '규제기관 일정 검색')
    except Exception as ex:
        print(f'  ⚠️  일정 병합 오류: {ex}')


if __name__ == '__main__':
    # 평일(월~금) 또는 INTL_CHECK=1 수동 실행 시 동작.
    # 각 소스는 5일 내 재검증을 건너뛰므로, 평일 실행해도 비용은 소스별 주 1회 수준 유지.
    is_weekday = datetime.now(KST).weekday() < 5   # 0=월 ~ 4=금
    forced = os.environ.get('INTL_CHECK') == '1'
    if is_weekday or forced:
        intl_check()
        check_schedules()
    else:
        print('⏭  주말 — 국제비교 검증 건너뜀 (INTL_CHECK=1 로 강제 실행 가능)')
